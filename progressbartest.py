def printProgressBar (iteration, total, prefix = '', suffix = '', decimals = 1, length = 100, fill = '█'):
   
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filledLength = int(length * iteration // total)
    bar = fill * filledLength + '-' * (length - filledLength)
    print('\r%s |%s| %s%% %s' % (prefix, bar, percent, suffix), end = '\r')
    # Print New Line on Complete
    if iteration == total: 
        print()

from time import sleep
from docx import Document

nb = input('Choose a number: ')
number = int(nb)
items = list(range(0, number))
l = len(items)

printProgressBar(0, l, prefix = 'Progress:', suffix = 'Complete', length = number)

for i in range(31, number):
    sleep(0.5)
    
    printProgressBar(i + 1, l, prefix = 'Progress:', suffix = 'Complete', length = number)
    document = Document('ebedjegy30.docx')
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '30' in paragraph.text:
                        paragraph.text = str(i)
                        document.save("ebedjegy" + str(i) + ".docx")
    