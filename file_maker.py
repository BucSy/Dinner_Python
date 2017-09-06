from time import sleep
from docx import Document
from print_progressbar import printProgressBar
import urllib.request
import os.path


if os.path.exists("ebedjegy.docx")==False:
    url = "https://www.dropbox.com/s/184r645dq4usohs/ebedjegy.docx?dl=1"
    urllib.request.urlretrieve(url, "ebedjegy.docx")
    sleep(2.6)
nb = input('How many Documents should I make?  ')
number = int(nb) + 1
items = list(range(0, number))
l = len(items)

printProgressBar(0, l, prefix = 'Progress:', suffix = 'Complete', length = number)

for i in range(1, number):
    sleep(0.5)
    
    printProgressBar(i + 1, l, prefix = 'Progress:', suffix = 'Complete', length = number)
    document = Document('ebedjegy.docx')
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'X' in paragraph.text:
                        paragraph.text = str(i)
                        document.save("ebedjegy" + str(i) + ".docx")
    